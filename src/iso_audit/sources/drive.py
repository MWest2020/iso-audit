"""Drive-source-adapter — Google Drive documenten lezen via `gws` CLI.

Implementeert het `Source` Protocol. Zoekt in de geconfigureerde Drive-map
(`AUDIT_SOURCE_FOLDER_ID` of `AUDIT_DRIVE_FOLDER_ID`) naar procedures,
werkinstructies en beleidsdocumenten.

Gemigreerd uit `Ops_to_Biz/audit/drive_ingest.py` per milestone B §2.3.2.
De legacy `haal_documenten_op()`-functie blijft als module-level callable
zodat bestaande callers (zoals de pipeline-CLI) ongewijzigd kunnen blijven
totdat ze omgezet zijn naar `SourceRegistry`-based dispatch.
"""

from __future__ import annotations

import io
import logging
import os
from collections.abc import Iterator
from typing import Any

import docx

from iso_audit.clients.gws import (
    gws_download_bestand,
    gws_exporteer_google_doc,
    gws_lijst_bestanden,
)
from iso_audit.sources import register
from iso_audit.sources.base import Document, Finding

logger = logging.getLogger(__name__)

BATCH_SIZE = 20

# Referentiedocumenten die geen organisatie-bewijs zijn — uitsluiten van
# classificatie. Begint-met-prefix-matching.
UITGESLOTEN_NAAM_PREFIXEN: tuple[str, ...] = (
    "NEN-EN-ISO",
    "ISO_IEC",
    "About the Sample Files",
)

ONDERSTEUNDE_MIME_TYPES: dict[str, str] = {
    "application/vnd.google-apps.document": "google_doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}
NIET_TEKSTUEEL: frozenset[str] = frozenset(
    {
        "application/pdf",
        "image/jpeg",
        "image/png",
        "image/gif",
        "image/tiff",
        "application/vnd.google-apps.presentation",
    }
)

FOLDER_ENV_VARS: tuple[str, ...] = ("AUDIT_SOURCE_FOLDER_ID", "AUDIT_DRIVE_FOLDER_ID")


def _split_ids(raw: str) -> list[str]:
    """Splits een komma-gescheiden string in losse, gestripte folder-ids."""
    return [v.split("?")[0].strip() for v in raw.split(",") if v.strip()]


def _resolve_folder_ids(expliciet: str | list[str] | None = None) -> list[str]:
    """Bepaal de Drive-folder-IDs (één of meer); raise als geen bron beschikbaar.

    Bronvolgorde:

    1. `expliciet` argument (string met komma-sep of list[str]).
    2. `AUDIT_SOURCE_FOLDER_ID` env (komma-sep) — primair voor multi-folder.
    3. `AUDIT_DRIVE_FOLDER_ID` env (komma-sep) — fallback.
    4. Beide env-vars samen: als ze allebei gezet en verschillend zijn, worden
       de IDs uit beide samengevoegd. Zo werken historische single-folder-
       configs én de multi-folder-praktijk waar Conduction zowel een Shared
       Drive (`0A...`) als een losse folder gebruikt.
    """
    if expliciet:
        if isinstance(expliciet, list):
            return [i for v in expliciet for i in _split_ids(v)]
        return _split_ids(expliciet)
    accumulated: list[str] = []
    for var in FOLDER_ENV_VARS:
        v = os.environ.get(var)
        if v:
            for fid in _split_ids(v):
                if fid not in accumulated:
                    accumulated.append(fid)
    if not accumulated:
        raise OSError(
            f"Geen Drive-map geconfigureerd. Stel {' of '.join(FOLDER_ENV_VARS)} in .env in."
        )
    return accumulated


def _resolve_folder_id(expliciet: str | None = None) -> str:
    """Backwards-compat: retourneer de eerste folder-ID uit de configuratie."""
    return _resolve_folder_ids(expliciet)[0]


def _is_uitgesloten(naam: str) -> bool:
    return any(naam.startswith(p) for p in UITGESLOTEN_NAAM_PREFIXEN)


def _fetch_tekst(file_id: str, mime: str) -> str:
    """Haal de tekst-inhoud op voor een Drive-bestand op basis van MIME."""
    if mime == "application/vnd.google-apps.document":
        return gws_exporteer_google_doc(file_id)
    inhoud = gws_download_bestand(file_id)
    if mime == ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        doc = docx.Document(io.BytesIO(inhoud))
        return "\n".join(p.text for p in doc.paragraphs)
    return inhoud.decode("utf-8", errors="replace")


@register
class DriveSource:
    """Google Drive `Source`-adapter via `gws` CLI."""

    naam = "drive"

    def __init__(self, folder_id: str | list[str] | None = None) -> None:
        """Configuratie wordt eenmalig vastgezet (immutable runtime-conf)."""
        self._folder_ids = _resolve_folder_ids(folder_id)
        # Shared Drive roots beginnen met "0A".
        self._drive_id_voor: dict[str, str | None] = {
            fid: (fid if fid.startswith("0A") else None) for fid in self._folder_ids
        }

    @property
    def folder_id(self) -> str:
        """Backwards-compat: eerste folder-ID."""
        return self._folder_ids[0]

    @property
    def folder_ids(self) -> list[str]:
        """Volledige lijst van geconfigureerde Drive-locaties."""
        return list(self._folder_ids)

    @property
    def drive_id(self) -> str | None:
        """Backwards-compat: drive-id van de eerste folder."""
        return self._drive_id_voor[self._folder_ids[0]]

    def list_documents(self, filter: dict[str, object] | None = None) -> Iterator[Document]:
        """Yield documenten uit alle geconfigureerde Drive-locaties (recursief).

        Bij meerdere folders wordt op `file-id` gededupliceerd zodat een
        bestand dat in twee scopes voorkomt maar één keer wordt opgeleverd.

        `filter` wordt momenteel genegeerd; Drive-filtering gebeurt op
        folder-niveau via env-configuratie (`AUDIT_SOURCE_FOLDER_ID`).
        Niet-tekstuele en onbekende MIME-types worden gelogd en geskipt.
        """
        del filter  # toekomstige uitbreiding; nu nog niet ondersteund
        logger.info(
            "DriveSource list_documents: folders=%s",
            self._folder_ids,
        )
        gezien: set[str] = set()
        for fid in self._folder_ids:
            for bestand in gws_lijst_bestanden(fid, drive_id=self._drive_id_voor[fid]):
                file_id = bestand["id"]
                if file_id in gezien:
                    continue
                gezien.add(file_id)
                naam = bestand["name"]
                mime = bestand["mimeType"]
                if _is_uitgesloten(naam):
                    logger.debug("Uitgesloten (referentiedocument): %s", naam)
                    continue
                if mime in NIET_TEKSTUEEL:
                    logger.info("Skip (niet-tekstueel): %s (%s)", naam, mime)
                    continue
                if mime not in ONDERSTEUNDE_MIME_TYPES:
                    logger.debug("Skip (onbekend MIME): %s (%s)", naam, mime)
                    continue
                yield Document(
                    id=file_id,
                    titel=naam,
                    bron="drive",
                    type=ONDERSTEUNDE_MIME_TYPES[mime],
                    laatst_gewijzigd=bestand.get("modifiedTime", ""),
                    inhoud_uri=file_id,
                )

    def fetch_content(self, doc: Document) -> str:
        """Lees de feitelijke tekst van een `Document` op uit Drive."""
        if doc.bron != self.naam:
            raise ValueError(
                f"DriveSource krijgt document uit bron={doc.bron!r}, verwacht {self.naam!r}"
            )
        mime_voor_type: dict[str, str] = {v: k for k, v in ONDERSTEUNDE_MIME_TYPES.items()}
        mime = mime_voor_type.get(doc.type)
        if not mime:
            raise ValueError(f"Onbekend Document-type voor DriveSource: {doc.type!r}")
        return _fetch_tekst(doc.inhoud_uri, mime)

    def list_findings(self, sessie_id: str) -> Iterator[Finding]:
        """Drive levert geen findings direct — een lege iterator."""
        del sessie_id
        return iter([])

    def probe(self) -> dict[str, object]:
        """Lichte connectiviteits-probe voor de UI grey-out (geen volledige listing).

        Eén bounded `files list` per folder (pageSize=1, niet-recursief) — bewijst
        auth + bereikbaarheid in een fractie van de tijd van ``healthcheck()``.
        """
        from iso_audit.clients.gws import gws_drive_bereikbaar

        for fid in self._folder_ids:
            try:
                gws_drive_bereikbaar(fid, drive_id=self._drive_id_voor[fid])
            except Exception as e:
                return {"status": "fail", "naam": self.naam, "tenant": fid, "reden": str(e)[:200]}
        return {"status": "ok", "naam": self.naam, "folders": list(self._folder_ids)}

    def healthcheck(self) -> dict[str, object]:
        """Verifieer dat alle geconfigureerde Drive-locaties bereikbaar zijn."""
        per_folder: dict[str, int] = {}
        for fid in self._folder_ids:
            try:
                bestanden = gws_lijst_bestanden(fid, drive_id=self._drive_id_voor[fid])
            except Exception as e:
                return {
                    "status": "fail",
                    "naam": self.naam,
                    "tenant": fid,
                    "reden": f"gws-fout op {fid}: {e}",
                }
            per_folder[fid] = len(bestanden)
        return {
            "status": "ok",
            "naam": self.naam,
            "folders": list(self._folder_ids),
            "per_folder": per_folder,
            "aantal_bestanden": sum(per_folder.values()),
        }


# ---------------------------------------------------------------------------
# Legacy API — pre-Source-protocol, blijft beschikbaar tot pipeline-cutover
# ---------------------------------------------------------------------------


def _verwerk_batch(
    batch: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    documenten: list[dict[str, Any]] = []
    handmatige_review: list[dict[str, Any]] = []
    for bestand in batch:
        naam = bestand["name"]
        file_id = bestand["id"]
        mime = bestand["mimeType"]
        if _is_uitgesloten(naam):
            logger.info("Uitgesloten (referentiedocument): %s", naam)
            continue
        if mime in NIET_TEKSTUEEL:
            handmatige_review.append(
                {
                    "naam": naam,
                    "id": file_id,
                    "reden": f"Niet-tekstueel formaat: {mime}",
                    "herkomst": "Drive",
                }
            )
            logger.info("Handmatige review vereist: %s (%s)", naam, mime)
            continue
        if mime not in ONDERSTEUNDE_MIME_TYPES:
            logger.debug("Onbekend mime-type overgeslagen: %s (%s)", naam, mime)
            continue
        try:
            tekst = _fetch_tekst(file_id, mime)
            documenten.append(
                {
                    "naam": naam,
                    "id": file_id,
                    "mime_type": mime,
                    "tekst": tekst,
                    "herkomst": "Drive",
                    "modified_at": bestand.get("modifiedTime"),
                }
            )
            logger.debug("Ingelezen: %s", naam)
        except Exception as e:
            logger.warning("Fout bij inlezen %s: %s", naam, e)
            handmatige_review.append(
                {
                    "naam": naam,
                    "id": file_id,
                    "reden": f"Leesfout: {e}",
                    "herkomst": "Drive",
                }
            )
    return documenten, handmatige_review


def haal_documenten_op(
    folder_id: str | list[str] | None = None,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Legacy-API: haal documenten op + lijst van items die handmatige review nodig hebben.

    Ondersteunt meerdere folders (komma-sep `AUDIT_SOURCE_FOLDER_ID` of
    `folder_id`-lijst); deduplicatie op file-id.

    Voor nieuwe code: gebruik `DriveSource.list_documents()` + `fetch_content()`.
    """
    resolved_ids = _resolve_folder_ids(folder_id)
    logger.info("Drive-ingest gestart vanuit %d locatie(s): %s", len(resolved_ids), resolved_ids)

    alle_bestanden: list[dict[str, Any]] = []
    gezien_ids: set[str] = set()
    for fid in resolved_ids:
        drive_id = fid if fid.startswith("0A") else None
        bestanden = gws_lijst_bestanden(fid, drive_id=drive_id)
        for b in bestanden:
            if b["id"] in gezien_ids:
                continue
            gezien_ids.add(b["id"])
            alle_bestanden.append(b)
        logger.info(
            "  %s (shared_drive=%s): %d bestanden (na dedup)",
            fid,
            bool(drive_id),
            len(alle_bestanden),
        )

    if not alle_bestanden:
        raise RuntimeError(
            f"Geen bestanden gevonden in Drive-locaties {resolved_ids}. "
            "Controleer de map-IDs en gws-authenticatie (`gws auth login`)."
        )
    logger.info("Totaal uniek gevonden: %d bestanden", len(alle_bestanden))

    alle_documenten: list[dict[str, Any]] = []
    alle_handmatige_review: list[dict[str, Any]] = []
    for i in range(0, len(alle_bestanden), BATCH_SIZE):
        batch = alle_bestanden[i : i + BATCH_SIZE]
        logger.info(
            "Verwerken batch %d/%d (%d bestanden)",
            i // BATCH_SIZE + 1,
            -(-len(alle_bestanden) // BATCH_SIZE),
            len(batch),
        )
        docs, review = _verwerk_batch(batch)
        alle_documenten.extend(docs)
        alle_handmatige_review.extend(review)

    logger.info(
        "Drive-ingest klaar: %d documenten ingelezen, %d voor handmatige review",
        len(alle_documenten),
        len(alle_handmatige_review),
    )
    return alle_documenten, alle_handmatige_review
